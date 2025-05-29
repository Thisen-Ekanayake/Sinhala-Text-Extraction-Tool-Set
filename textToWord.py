from tkinter import Tk, filedialog
from docx import Document
from pathlib import Path

def merge_txt_to_docx():
    # Hide the root Tkinter window
    root = Tk()
    root.withdraw()

    # File picker to select multiple .txt files
    file_paths = filedialog.askopenfilenames(
        title="Select text files to merge",
        filetypes=[("Text Files", "*.txt")]
    )

    if not file_paths:
        print("No files selected.")
        return

    # Create a new Word document
    doc = Document()

    for file_path in file_paths:
        path = Path(file_path)
        with open(path, "r", encoding="utf-8") as file:
            content = file.read()

        # Add file name as a heading
        doc.add_heading(path.name, level=2)
        doc.add_paragraph(content)
        doc.add_page_break()  # Optional: adds page break between files

        print(f"Added: {path.name}")

    # Save output file
    output_path = filedialog.asksaveasfilename(
        title="Save Merged DOCX As",
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")]
    )

    if output_path:
        doc.save(output_path)
        print(f"\nMerged DOCX saved as: {output_path}")
    else:
        print("\nSave cancelled.")

if __name__ == "__main__":
    merge_txt_to_docx()
