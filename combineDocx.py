from tkinter import Tk, filedialog
from docx import Document

def get_word_count(doc):
    count = 0
    for para in doc.paragraphs:
        count += len(para.text.strip().split())
    return count

def combine_and_count():
    root = Tk()
    root.withdraw()

    # Select the first DOCX file
    file1 = filedialog.askopenfilename(title="Select first DOCX file", filetypes=[("Word Documents", "*.docx")])
    if not file1:
        print("First file not selected. Exiting.")
        return

    # Select the second DOCX file
    file2 = filedialog.askopenfilename(title="Select second DOCX file", filetypes=[("Word Documents", "*.docx")])
    if not file2:
        print("Second file not selected. Exiting.")
        return

    # Load documents
    doc1 = Document(file1)
    doc2 = Document(file2)

    # Combine doc2 into doc1
    for para in doc2.paragraphs:
        doc1.add_paragraph(para.text)

    # Calculate total word count
    total_words = get_word_count(doc1)

    print(f"\nFiles combined:\n 1. {file1}\n 2. {file2}")
    print(f"Total word count: {total_words}")

if __name__ == "__main__":
    combine_and_count()
